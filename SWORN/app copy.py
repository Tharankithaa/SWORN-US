import os
import re
import cv2
import numpy as np
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
import streamlit as st
import base64
import tempfile
from io import BytesIO
from googletrans import Translator

# ----------------------------
# Configuration and Setup
# ----------------------------

# Website Configuration
st.set_page_config(page_title="SWORN", layout="wide")

# Set the background color
st.markdown(
    """
    <style>
        .stApp {
            background-color: #ffd9c2;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# Path for logo
logo_path = r"C:\Users\HP\Downloads\SWORN\logo.png"  # Use the full path to your logo file

# Function to encode the logo in base64
def get_base64_image(path):
    with open(path, "rb") as file:
        return base64.b64encode(file.read()).decode("utf-8")

# Display logo at the center
try:
    logo_base64 = get_base64_image(logo_path)
    st.markdown(
        f"""
        <div style='text-align: center;'>
            <img src='data:image/png;base64,{logo_base64}' alt='Logo' style='width: 150px;'>
        </div>
        """,
        unsafe_allow_html=True,
    )
except Exception as e:
    st.error("Logo not found. Please ensure the correct path is specified.")

# Centered Welcome Title
st.markdown(
    """
    <h1 style='text-align: center; color: BLACK;'>Welcome to SWORN!</h1>
    """,
    unsafe_allow_html=True,
)

st.write("Upload an image or PDF document to extract text!")

# ----------------------------
# Helper Functions
# ----------------------------

# Text cleaning function
def clean_text(text):
    # Remove HTML tags
    text = re.sub(r'<.*?>', '', text)
    # Remove special characters
    text = re.sub(r'[@#$%^&*_+=<>~|]', '', text)
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text

# Function to perform OCR on an image
def perform_ocr(image):
    # Preprocessing steps
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)  # Convert to grayscale
    normalized = cv2.normalize(gray, None, 0, 255, cv2.NORM_MINMAX)  # Normalize brightness/contrast
    blurred = cv2.medianBlur(normalized, 3)  # Reduce noise
    sharpen_kernel = np.array([[-1, -1, -1],
                               [-1, 9, -1],
                               [-1, -1, -1]])  # Sharpen kernel
    sharpened = cv2.filter2D(blurred, -1, sharpen_kernel)
    _, binary = cv2.threshold(sharpened, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)  # Otsu's thresholding

    # Perform OCR on the preprocessed image
    languages = "eng+tam+sin+Sinhala+Tamil"  # Specify languages for OCR
    text = pytesseract.image_to_string(binary, lang=languages, config='--psm 4')

    # Clean extracted text
    cleaned_text = clean_text(text)
    return cleaned_text

# Function to create a Word document from extracted text
def create_word_document(extracted_text, is_pdf=False):
    doc = Document()
    if is_pdf:
        doc.add_heading('Extracted Text from PDF', level=1)
        for page_num, text in enumerate(extracted_text, start=1):
            doc.add_heading(f'Page {page_num}', level=2)
            doc.add_paragraph(text)
    else:
        doc.add_heading('Extracted Text from Image', level=1)
        doc.add_paragraph(extracted_text)

    # Save the document to a BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Function to translate text
def translate_text(text, target_language):
    translator = Translator()
    translated = translator.translate(text, dest=target_language)
    return translated.text

# ----------------------------
# File Upload and Processing
# ----------------------------

# Upload Section
uploaded_file = st.file_uploader("Upload Image or Document", type=["png", "jpg", "jpeg", "pdf"])

# Define extracted_text as None initially
extracted_text = None

if uploaded_file is not None:
    file_type = uploaded_file.type
    file_name = uploaded_file.name

    if file_type in ["image/png", "image/jpeg"]:
        # ----------------------------
        # Image Processing
        # ----------------------------
        st.subheader("Image Uploaded: " + file_name)
        
        # Load the image
        image = Image.open(uploaded_file)
        image_np = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

        # Perform OCR
        extracted_text = perform_ocr(image_np)

        # Display the extracted text
        st.text_area("Extracted Text", extracted_text, height=300)

        # Create and provide a download link for the Word document
        if extracted_text:
            doc_io = create_word_document(extracted_text, is_pdf=False)
            st.download_button(
                label="Download Extracted Text as Word Document",
                data=doc_io,
                file_name="extracted_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif file_type == "application/pdf":
        # ----------------------------
        # PDF Processing
        # ----------------------------
        st.subheader("PDF Uploaded: " + file_name)
        
        # Save the uploaded PDF to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(uploaded_file.getbuffer())
            temp_pdf_path = temp_pdf.name

        try:
            # Convert PDF pages to images
            pdf_pages = convert_from_path(temp_pdf_path, dpi=300)  # 300 dpi for high-quality images

            extracted_text_pages = []
            for page_number, page_image in enumerate(pdf_pages, start=1):
                st.write(f"Processing Page {page_number}...")

                # Convert PIL image to OpenCV format
                image = np.array(page_image)
                image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)

                # Perform OCR on the image
                text = perform_ocr(image)

                # Ensure text is a string
                if not isinstance(text, str):
                    text = str(text)

                extracted_text_pages.append(text)

            # Display the extracted text
            combined_text = "\n\n".join([f"Page {i+1}:\n{text}" for i, text in enumerate(extracted_text_pages)])
            st.text_area("Extracted Text from PDF", combined_text, height=300)

            # Create and provide a download link for the Word document
            if any(extracted_text_pages):
                doc_io = create_word_document(extracted_text_pages, is_pdf=True)
                st.download_button(
                    label="Download Extracted Text as Word Document",
                    data=doc_io,
                    file_name="extracted_text.pdf.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Error processing PDF: {e}")

        finally:
            # Clean up the temporary PDF file
            os.remove(temp_pdf_path)

    else:
        st.error("Unsupported file type. Please upload a valid image or PDF file.")

# ----------------------------
# Language Selection for Translation
# ----------------------------

